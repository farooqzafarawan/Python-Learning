from docx import Document
import pandas as pd
from typing import Dict, List, Tuple, Optional

class NestedTableExtractor:
    def __init__(self, docx_path: str):
        """
        Initialize the extractor with document path.
        
        Args:
            docx_path (str): Path to the Word document
        """
        self.doc = Document(docx_path)
        self.reserves_tables = []  # Store nested tables
        self.main_tables = []      # Store main tables
        
    def extract_cell_content(self, cell) -> Tuple[str, Optional[pd.DataFrame], str]:
        """
        Extract content from a cell, separating text before table, nested table, and text after.
        
        Args:
            cell: Word document table cell
            
        Returns:
            Tuple containing:
            - Text before table
            - DataFrame of nested table (if found)
            - Text after table
        """
        before_text = []
        after_text = []
        current_section = before_text
        found_table = False
        nested_table_data = None
        
        # Process each paragraph and table in the cell
        for element in cell._element:
            if element.tag.endswith('p'):  # Paragraph
                paragraph_text = element.text.strip()
                if paragraph_text:
                    current_section.append(paragraph_text)
                    
            elif element.tag.endswith('tbl'):  # Table
                found_table = True
                current_section = after_text
                
                # Convert nested table element to Table object
                nested_table = cell.tables[0]  # Assuming it's the first nested table
                
                # Extract nested table data
                if "CURRENT & INCREASED RESERVES" in ' '.join(before_text):
                    nested_data = []
                    for row in nested_table.rows:
                        cell_text = row.cells[0].text.strip()  # Single column table
                        if cell_text.startswith('Reinsurance:'):
                            nested_data.append([cell_text])
                        else:
                            nested_data.append([cell_text])
                    
                    # Create DataFrame
                    nested_table_data = pd.DataFrame(nested_data, columns=['Reinsurance Details'])
                    # Store in reserves_tables list
                    self.reserves_tables.append(nested_table_data)
        
        return (
            '\n'.join(before_text),
            nested_table_data,
            '\n'.join(after_text)
        )
    
    def process_document(self) -> Dict[str, List[Dict]]:
        """
        Process the entire document, extracting both main and nested tables.
        
        Returns:
            Dict containing processed tables and their context
        """
        results = {
            'main_tables': [],
            'nested_reserves_tables': []
        }
        
        # Process each table in the document
        for table_idx, table in enumerate(self.doc.tables):
            table_data = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = {}
                
                for cell_idx, cell in enumerate(row.cells):
                    # Extract cell content with potential nested table
                    before_text, nested_table, after_text = self.extract_cell_content(cell)
                    
                    cell_info = {
                        'position': f'R{row_idx+1}C{cell_idx+1}',
                        'before_text': before_text,
                        'after_text': after_text,
                        'has_nested_table': nested_table is not None
                    }
                    
                    row_data[f'cell_{cell_idx+1}'] = cell_info
                
                table_data.append(row_data)
            
            results['main_tables'].append({
                'table_index': table_idx + 1,
                'data': table_data
            })
        
        # Add nested reserves tables
        results['nested_reserves_tables'] = [
            df.to_dict('records') for df in self.reserves_tables
        ]
        
        return results

def save_results(results: Dict[str, List[Dict]], output_base_path: str):
    """
    Save extracted data to separate files.
    
    Args:
        results: Dictionary containing extracted data
        output_base_path: Base path for output files
    """
    # Save main tables context
    with open(f"{output_base_path}_main_tables.txt", 'w', encoding='utf-8') as f:
        for table in results['main_tables']:
            f.write(f"\nTable {table['table_index']}:\n")
            f.write("="*50 + "\n")
            
            for row in table['data']:
                for cell_key, cell_data in row.items():
                    f.write(f"\n{cell_data['position']}:\n")
                    if cell_data['before_text']:
                        f.write(f"Text before table:\n{cell_data['before_text']}\n")
                    if cell_data['has_nested_table']:
                        f.write("Contains nested reserves table\n")
                    if cell_data['after_text']:
                        f.write(f"Text after table:\n{cell_data['after_text']}\n")
                    f.write("-"*40 + "\n")
    
    # Save nested reserves tables to Excel
    if results['nested_reserves_tables']:
        df_list = [pd.DataFrame(table) for table in results['nested_reserves_tables']]
        with pd.ExcelWriter(f"{output_base_path}_reserves_tables.xlsx") as writer:
            for idx, df in enumerate(df_list):
                df.to_excel(writer, sheet_name=f'Reserves_Table_{idx+1}', index=False)

# Example usage
if __name__ == "__main__":
    docx_path = "sample.docx"
    output_base_path = "extracted_data"
    
    try:
        # Initialize extractor
        extractor = NestedTableExtractor(docx_path)
        
        # Process document
        results = extractor.process_document()
        
        # Print summary
        print(f"Found {len(results['main_tables'])} main tables")
        print(f"Found {len(results['nested_reserves_tables'])} reserves tables")
        
        # Save results
        save_results(results, output_base_path)
        print(f"\nResults saved with base path: {output_base_path}")
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")