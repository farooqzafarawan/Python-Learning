from docx import Document
from typing import Dict, Optional
import re

def get_formatted_text_runs(paragraph) -> list:
    """
    Extract text runs with their formatting from a paragraph.
    
    Args:
        paragraph: A paragraph object from python-docx
        
    Returns:
        list: List of tuples containing (text, is_bold) for each run
    """
    runs = []
    for run in paragraph.runs:
        # Get text and remove extra whitespace
        text = run.text.strip()
        if text:
            runs.append((text, run.bold or False))
    return runs

def process_cell_content(cell) -> Dict[str, Optional[str]]:
    """
    Process cell content to extract dictionary items based on bold text patterns.
    
    Args:
        cell: A table cell object from python-docx
        
    Returns:
        Dict[str, Optional[str]]: Dictionary with bold text as keys and following text as values
    """
    result_dict = {}
    current_key = None
    current_value = []
    buffer = []
    
    # Process each paragraph in the cell
    for paragraph in cell.paragraphs:
        runs = get_formatted_text_runs(paragraph)
        
        # Skip empty paragraphs
        if not runs:
            if current_key and current_value:
                # Add a line break if we're collecting a value
                current_value.append('\n')
            continue
        
        for text, is_bold in runs:
            # If we find bold text ending with colon
            if is_bold and ':' in text:
                # If we have a previous key-value pair, save it
                if current_key:
                    value_text = ' '.join(current_value).strip()
                    result_dict[current_key] = value_text if value_text else None
                    current_value = []
                
                # Extract new key (bold text up to the colon)
                key_match = re.match(r'(.*?):', text)
                if key_match:
                    current_key = key_match.group(1).strip()
                    # Add any remaining text after colon to value
                    remaining_text = text[len(current_key) + 1:].strip()
                    if remaining_text:
                        current_value.append(remaining_text)
            
            # If we find non-bold text ending with colon (no value case)
            elif ':' in text and not is_bold:
                # Save previous key-value pair if exists
                if current_key:
                    value_text = ' '.join(current_value).strip()
                    result_dict[current_key] = value_text if value_text else None
                
                # Extract new key (text up to the colon)
                key_match = re.match(r'(.*?):', text)
                if key_match:
                    key = key_match.group(1).strip()
                    result_dict[key] = None
                    current_key = None
                    current_value = []
            
            # Regular text (potential value)
            elif current_key:
                current_value.append(text)
            
            # Non-bold text without colon (collect in buffer)
            else:
                buffer.append(text)
    
    # Save last key-value pair if exists
    if current_key:
        value_text = ' '.join(current_value).strip()
        result_dict[current_key] = value_text if value_text else None
    
    # Process any remaining buffer text
    if buffer:
        buffer_text = ' '.join(buffer).strip()
        if buffer_text:
            result_dict[buffer_text] = None
    
    return result_dict

def process_word_document(docx_path: str) -> Dict[str, Dict[str, Optional[str]]]:
    """
    Process a Word document and extract formatted text patterns from all tables.
    
    Args:
        docx_path (str): Path to the Word document
        
    Returns:
        Dict[str, Dict[str, Optional[str]]]: Dictionary containing table data
    """
    doc = Document(docx_path)
    document_data = {}
    
    for table_index, table in enumerate(doc.tables):
        table_dict = {}
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_dict = process_cell_content(cell)
                if cell_dict:
                    cell_key = f"Cell_{row_index}_{cell_index}"
                    table_dict[cell_key] = cell_dict
        
        if table_dict:
            document_data[f"Table_{table_index}"] = table_dict
    
    return document_data

def save_to_text_file(data: Dict[str, Dict[str, Optional[str]]], output_path: str):
    """
    Save the extracted data to a text file in a readable format.
    
    Args:
        data: Dictionary containing the extracted data
        output_path (str): Path for the output text file
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        for table_name, table_data in data.items():
            f.write(f"\n{table_name}:\n{'='*50}\n")
            for cell_name, cell_data in table_data.items():
                f.write(f"\n{cell_name}:\n{'-'*30}\n")
                for key, value in cell_data.items():
                    if value is None:
                        f.write(f"• {key}: <No Value>\n")
                    else:
                        f.write(f"• {key}: {value}\n")
                f.write('\n')

# Example usage
if __name__ == "__main__":
    docx_path = "sample.docx"
    output_path = "extracted_patterns.txt"
    
    try:
        # Process the document
        extracted_data = process_word_document(docx_path)
        
        # Print summary
        total_items = sum(
            len(cell_data) 
            for table in extracted_data.values() 
            for cell_data in table.values()
        )
        print(f"Successfully processed document")
        print(f"Found {len(extracted_data)} tables")
        print(f"Extracted {total_items} total key-value pairs")
        
        # Save results
        save_to_text_file(extracted_data, output_path)
        print(f"\nResults saved to {output_path}")
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")